import json
import os
from docx import Document
from docx.shared import Inches
import PIL.Image
import imagehash

def load_json_data(filepath):
    """Utility function to load JSON data from a file."""
    with open(filepath, 'r') as file:
        return json.load(file)

def find_screenshots_for_clip(clip_number, screenshots_folder):
    """Find all screenshots for a given clip number in the specified folder."""
    prefix = f"clip_{clip_number}_"
    return [f for f in os.listdir(screenshots_folder) if f.startswith(prefix)]

def create_word_document(video_summary_file, screenshots_folder, output_file):
    summaries = load_json_data(video_summary_file)
    doc = Document()
    doc.add_heading('Video Clip Summaries and Screenshots', level=1)

    # Iterate through each video clip description
    for clip_id, description in summaries.items():
        doc.add_heading(f"Step {clip_id.split('clip')[1]}", level=1)
        doc.add_paragraph(description)

        clip_number = clip_id.split('clip')[1]
        screenshot_filenames = find_screenshots_for_clip(clip_number, screenshots_folder)

        added_hashes = []  # List to store hashes of added images
        threshold = 2  # Difference threshold for image similarity (lower means more similar)

        # Add screenshots if available and not similar to already added ones
        for filename in screenshot_filenames:
            path = os.path.join(screenshots_folder, filename)
            if os.path.exists(path):
                try:
                    # Compute the hash of the image
                    img = PIL.Image.open(path)
                    img_hash = imagehash.average_hash(img)

                    # Check for similarity with already added images
                    if not any((img_hash - h) < threshold for h in added_hashes):
                        added_hashes.append(img_hash)
                        doc.add_picture(path, width=Inches(5.0), height=Inches(3.5))
                    else:
                        print(f"Similar image already added for {path}")
                except Exception as e:
                    print(f"Error adding picture {path}: {e}")
            else:
                print(f"Screenshot not found for {path}")

    # Save the document
    doc.save(output_file)

if __name__ == "__main__":
    # Define the paths to your JSON file and the output document
    video_summary_file = "Final_combined.json"
    screenshots_folder = "Screenshots"
    output_file = "Video_Summary.docx"

    # Call the function to create the Word document
    create_word_document(video_summary_file, screenshots_folder, output_file)
